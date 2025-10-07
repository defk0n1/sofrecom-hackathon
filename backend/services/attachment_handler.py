import os
import base64
import mimetypes
from typing import List, Dict, Any, Optional
import io

# For document processing
try:
    from PyPDF2 import PdfReader
    import docx
    from PIL import Image
    import pandas as pd
    import json
    import csv
except ImportError:
    print("Some libraries not installed. Install with: pip install PyPDF2 python-docx Pillow pandas openpyxl")


class AttachmentProcessor:
    """Process different file types for LLM analysis"""
    
    @staticmethod
    def process_pdf(file_data: bytes) -> Dict[str, Any]:
        """Extract text from PDF"""
        try:
            pdf_file = io.BytesIO(file_data)
            pdf_reader = PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                text_content.append(f"--- Page {page_num + 1} ---\n{page.extract_text()}")
            
            return {
                'type': 'pdf',
                'num_pages': len(pdf_reader.pages),
                'text': '\n\n'.join(text_content),
                'metadata': pdf_reader.metadata if hasattr(pdf_reader, 'metadata') else {}
            }
        except Exception as e:
            return {'type': 'pdf', 'error': str(e), 'text': ''}

    @staticmethod
    def process_docx(file_data: bytes) -> Dict[str, Any]:
        """Extract text from DOCX"""
        try:
            doc_file = io.BytesIO(file_data)
            doc = docx.Document(doc_file)
            
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables_data = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    table_text.append([cell.text for cell in row.cells])
                tables_data.append(table_text)
            
            return {
                'type': 'docx',
                'text': '\n'.join(paragraphs),
                'tables': tables_data,
                'num_paragraphs': len(paragraphs),
                'num_tables': len(tables_data)
            }
        except Exception as e:
            return {'type': 'docx', 'error': str(e), 'text': ''}

    @staticmethod
    def process_text(file_data: bytes, encoding: str = 'utf-8') -> Dict[str, Any]:
        """Process plain text files"""
        try:
            text = file_data.decode(encoding)
            return {
                'type': 'text',
                'text': text,
                'line_count': len(text.split('\n')),
                'char_count': len(text)
            }
        except UnicodeDecodeError:
            # Try different encodings
            for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = file_data.decode(enc)
                    return {
                        'type': 'text',
                        'text': text,
                        'encoding': enc,
                        'line_count': len(text.split('\n')),
                        'char_count': len(text)
                    }
                except:
                    continue
            return {'type': 'text', 'error': 'Unable to decode text', 'text': ''}

    @staticmethod
    def process_csv(file_data: bytes) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            text_data = file_data.decode('utf-8')
            csv_file = io.StringIO(text_data)
            reader = csv.reader(csv_file)
            rows = list(reader)
            
            # Convert to structured format
            if rows:
                headers = rows[0]
                data = rows[1:]
                
                # Create summary
                summary = f"CSV with {len(headers)} columns and {len(data)} rows\n"
                summary += f"Columns: {', '.join(headers)}\n\n"
                summary += "Preview (first 5 rows):\n"
                for row in data[:5]:
                    summary += ' | '.join(str(cell) for cell in row) + '\n'
                
                return {
                    'type': 'csv',
                    'text': summary,
                    'headers': headers,
                    'row_count': len(data),
                    'data': data[:100]  # Limit data for LLM
                }
            return {'type': 'csv', 'text': 'Empty CSV file'}
        except Exception as e:
            return {'type': 'csv', 'error': str(e), 'text': ''}

    @staticmethod
    def process_excel(file_data: bytes) -> Dict[str, Any]:
        """Process Excel files"""
        try:
            excel_file = io.BytesIO(file_data)
            df_dict = pd.read_excel(excel_file, sheet_name=None)
            
            summary = f"Excel file with {len(df_dict)} sheet(s)\n\n"
            
            for sheet_name, df in df_dict.items():
                summary += f"--- Sheet: {sheet_name} ---\n"
                summary += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                summary += f"Columns: {', '.join(df.columns)}\n"
                summary += f"\nPreview:\n{df.head().to_string()}\n\n"
            
            return {
                'type': 'excel',
                'text': summary,
                'sheets': list(df_dict.keys()),
                'sheet_count': len(df_dict)
            }
        except Exception as e:
            return {'type': 'excel', 'error': str(e), 'text': ''}

    @staticmethod
    def process_json(file_data: bytes) -> Dict[str, Any]:
        """Process JSON files"""
        try:
            json_data = json.loads(file_data.decode('utf-8'))
            
            # Create readable summary
            summary = json.dumps(json_data, indent=2)
            
            return {
                'type': 'json',
                'text': summary,
                'data': json_data,
                'keys': list(json_data.keys()) if isinstance(json_data, dict) else None
            }
        except Exception as e:
            return {'type': 'json', 'error': str(e), 'text': ''}

    @staticmethod
    def process_image(file_data: bytes, filename: str) -> Dict[str, Any]:
        """Process image files - extract metadata"""
        try:
            image = Image.open(io.BytesIO(file_data))
            
            # Note: For actual OCR, you'd use pytesseract
            # import pytesseract
            # text = pytesseract.image_to_string(image)
            
            return {
                'type': 'image',
                'text': f"Image file: {filename}\nDimensions: {image.size[0]}x{image.size[1]}\nFormat: {image.format}\nMode: {image.mode}",
                'width': image.size[0],
                'height': image.size[1],
                'format': image.format,
                'mode': image.mode,
                'note': 'For OCR text extraction, install pytesseract'
            }
        except Exception as e:
            return {'type': 'image', 'error': str(e), 'text': ''}

    @staticmethod
    def process_html(file_data: bytes) -> Dict[str, Any]:
        """Process HTML files"""
        try:
            try:
                from bs4 import BeautifulSoup
                html_content = file_data.decode('utf-8')
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Extract text content
                text = soup.get_text(separator='\n', strip=True)
                
                return {
                    'type': 'html',
                    'text': text,
                    'title': soup.title.string if soup.title else None,
                    'links': [a.get('href') for a in soup.find_all('a', href=True)][:50]
                }
            except ImportError:
                # Fallback without BeautifulSoup
                html_content = file_data.decode('utf-8')
                return {
                    'type': 'html',
                    'text': html_content,
                    'note': 'Install beautifulsoup4 for better HTML parsing'
                }
        except Exception as e:
            return {'type': 'html', 'error': str(e), 'text': file_data.decode('utf-8', errors='ignore')}

    @staticmethod
    def process_code(file_data: bytes, extension: str) -> Dict[str, Any]:
        """Process code files"""
        try:
            code = file_data.decode('utf-8')
            lines = code.split('\n')
            
            return {
                'type': 'code',
                'language': extension.lstrip('.'),
                'text': code,
                'line_count': len(lines),
                'char_count': len(code)
            }
        except Exception as e:
            return {'type': 'code', 'error': str(e), 'text': ''}

    @staticmethod
    def process_file(file_data: bytes, filename: str) -> Dict[str, Any]:
        """Main entry point to process any file type"""
        file_ext = os.path.splitext(filename)[1].lower()
        
        processed_data = {
            'filename': filename,
            'size': len(file_data),
            'extension': file_ext
        }
        
        # PDF files
        if file_ext == '.pdf':
            processed_data.update(AttachmentProcessor.process_pdf(file_data))
        
        # Word documents
        elif file_ext in ['.docx', '.doc']:
            processed_data.update(AttachmentProcessor.process_docx(file_data))
        
        # Text files
        elif file_ext in ['.txt', '.md', '.log', '.rtf']:
            processed_data.update(AttachmentProcessor.process_text(file_data))
        
        # CSV files
        elif file_ext == '.csv':
            processed_data.update(AttachmentProcessor.process_csv(file_data))
        
        # Excel files
        elif file_ext in ['.xlsx', '.xls']:
            processed_data.update(AttachmentProcessor.process_excel(file_data))
        
        # JSON files
        elif file_ext == '.json':
            processed_data.update(AttachmentProcessor.process_json(file_data))
        
        # Image files
        elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
            processed_data.update(AttachmentProcessor.process_image(file_data, filename))
        
        # HTML files
        elif file_ext in ['.html', '.htm']:
            processed_data.update(AttachmentProcessor.process_html(file_data))
        
        # Code files
        elif file_ext in ['.py', '.js', '.java', '.cpp', '.c', '.cs', '.rb', '.go', '.rs', '.php', '.swift', '.kt', '.tsx', '.jsx']:
            processed_data.update(AttachmentProcessor.process_code(file_data, file_ext))
        
        # Unknown file type
        else:
            processed_data.update({
                'type': 'unknown',
                'text': f'Binary file: {filename} ({len(file_data)} bytes)',
                'note': 'Unsupported file type for text extraction'
            })
        
        return processed_data


class AttachmentFormatter:
    """Format processed attachments for LLM consumption"""
    
    @staticmethod
    def format_for_llm(
        processed_attachments: List[Dict[str, Any]],
        include_full_text: bool = True,
        max_text_length: int = 10000
    ) -> str:
        """Format processed attachments for LLM consumption"""
        
        if not processed_attachments:
            return "No attachments found."
        
        llm_text = f"Email contains {len(processed_attachments)} attachment(s):\n\n"
        
        for i, attachment in enumerate(processed_attachments, 1):
            llm_text += f"{'='*60}\n"
            llm_text += f"ATTACHMENT {i}: {attachment['filename']}\n"
            llm_text += f"Type: {attachment.get('type', 'unknown')}\n"
            llm_text += f"Size: {attachment['size']} bytes\n"
            
            # Add type-specific metadata
            if attachment.get('type') == 'pdf':
                llm_text += f"Pages: {attachment.get('num_pages', 'N/A')}\n"
            elif attachment.get('type') == 'excel':
                llm_text += f"Sheets: {', '.join(attachment.get('sheets', []))}\n"
            elif attachment.get('type') == 'image':
                llm_text += f"Dimensions: {attachment.get('width')}x{attachment.get('height')}\n"
            elif attachment.get('type') == 'csv':
                llm_text += f"Rows: {attachment.get('row_count', 'N/A')}\n"
            elif attachment.get('type') == 'docx':
                llm_text += f"Paragraphs: {attachment.get('num_paragraphs', 'N/A')}\n"
            
            llm_text += f"\n"
            
            # Add content
            if include_full_text and 'text' in attachment:
                text_content = attachment['text']
                if len(text_content) > max_text_length:
                    llm_text += f"CONTENT (truncated to {max_text_length} chars):\n"
                    llm_text += text_content[:max_text_length] + "\n... [truncated]\n"
                else:
                    llm_text += f"CONTENT:\n{text_content}\n"
            
            if 'error' in attachment:
                llm_text += f"ERROR: {attachment['error']}\n"
            
            llm_text += f"\n"
        
        return llm_text
    
    @staticmethod
    def format_summary(processed_attachments: List[Dict[str, Any]]) -> str:
        """Create a brief summary of attachments"""
        if not processed_attachments:
            return "No attachments"
        
        summary_parts = []
        for att in processed_attachments:
            file_type = att.get('type', 'unknown')
            filename = att.get('filename', 'unknown')
            summary_parts.append(f"{filename} ({file_type})")
        
        return f"{len(processed_attachments)} attachment(s): {', '.join(summary_parts)}"